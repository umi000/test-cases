from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_header(cell):
    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run("")
    run.font.bold = True
    run.font.size = Pt(10)


def set_table_style(table):
    # Set a light grid table style if available
    try:
        table.style = "Light Grid"
    except Exception:
        pass


def add_key_value(p, key, value):
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    p.add_run(value if value is not None else "")


def add_section_heading(document, text):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_test_case(document, tc):
    # Heading separator
    document.add_paragraph("\n")

    # Key fields
    add_key_value(document.add_paragraph(), "Test Case ID", tc["id"])  # ID
    add_key_value(document.add_paragraph(), "Test Objective", tc["objective"])  # Objective
    add_key_value(document.add_paragraph(), "Description", tc["description"])  # Description
    add_key_value(document.add_paragraph(), "Module", tc.get("module", "Client Portal Configuration"))
    add_key_value(document.add_paragraph(), "Priority", tc.get("priority", "P2"))
    add_key_value(document.add_paragraph(), "Test Type", tc.get("type", "Functional"))
    add_key_value(document.add_paragraph(), "Date", tc.get("date", "2025-10-18"))
    add_key_value(document.add_paragraph(), "Tester", tc.get("tester", "Umair"))
    add_key_value(document.add_paragraph(), "JIRA issue #", tc.get("jira", "[To be filled]"))
    add_key_value(document.add_paragraph(), "Precondition", tc.get("precondition", ""))

    # Steps title
    add_section_heading(document, "\nTest Steps:")

    # Steps table
    table = document.add_table(rows=1, cols=4)
    set_table_style(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Step #"
    hdr_cells[1].text = "User Action"
    hdr_cells[2].text = "Expected Result"
    hdr_cells[3].text = "Actual Result"
    for c in hdr_cells:
        set_cell_header(c)

    for idx, step in enumerate(tc["steps"], start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(step.get("step", idx))
        row_cells[1].text = step.get("action", "")
        row_cells[2].text = step.get("expected", "")
        row_cells[3].text = step.get("actual", "[Tester to fill after execution]")

    # Post condition
    add_key_value(document.add_paragraph(), "\nPost Condition", tc.get("postcondition", ""))


def build_document(output_path: str):
    document = Document()

    # Title
    title_p = document.add_paragraph()
    title_run = title_p.add_run("Client Portal Configuration - View Subscreen Test Cases")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.add_run("Generated on: 2025-10-18").italic = True

    default_meta = {
        "module": "Client Portal Configuration",
        "date": "2025-10-18",
        "tester": "Umair",
        "jira": "[To be filled]",
    }

    # Define all test cases
    tcs = [
        {
            "id": "TC_ClientPortalConfig_View_001",
            "objective": "Open View subscreen for a selected configuration",
            "description": "Verify that clicking View on a single selected configuration opens the read-only subscreen.",
            "priority": "P1",
            "type": "Functional",
            "precondition": "User logged in with access; at least one configuration exists.",
            "steps": [
                {"action": "Select a single configuration row (e.g., TIS) in the list", "expected": "Row becomes selected"},
                {"action": "Click View", "expected": "Read-only Client Portal Configuration subscreen opens for the selected record"},
            ],
            "postcondition": "View subscreen is open; no data modified.",
        },
        {
            "id": "TC_ClientPortalConfig_View_002",
            "objective": "Prevent View when no record is selected",
            "description": "Ensure the system requires a selection before opening View.",
            "priority": "P1",
            "type": "Negative",
            "precondition": "User is on the configuration list.",
            "steps": [
                {"action": "Do not select any configuration", "expected": "No row selected"},
                {"action": "Click View", "expected": "System shows a validation message and does not open the subscreen"},
            ],
            "postcondition": "No subscreen opened.",
        },
        {
            "id": "TC_ClientPortalConfig_View_003",
            "objective": "Prevent View with multiple selections",
            "description": "Verify that View requires exactly one selection.",
            "priority": "P1",
            "type": "Negative",
            "precondition": "List supports multi-select.",
            "steps": [
                {"action": "Select two or more configuration rows", "expected": "Multiple rows highlighted"},
                {"action": "Click View", "expected": "System prompts to select only one record and does not open the subscreen"},
            ],
            "postcondition": "No subscreen opened.",
        },
        {
            "id": "TC_ClientPortalConfig_View_004",
            "objective": "Enforce read-only mode in View",
            "description": "Ensure all fields and lists are non-editable in the View subscreen.",
            "priority": "P1",
            "type": "Functional",
            "precondition": "View subscreen is open for any configuration.",
            "steps": [
                {"action": "Attempt to type in Configuration Name", "expected": "Field is read-only; no changes allowed"},
                {"action": "Try toggling any checkbox (e.g., Active)", "expected": "Checkbox cannot be changed"},
                {"action": "Try adding/removing items from any list (Reports/Clients/Account Managers)", "expected": "No add/remove controls available; lists are not editable"},
            ],
            "postcondition": "No data modified.",
        },
        {
            "id": "TC_ClientPortalConfig_View_005",
            "objective": "Default landing tab on open",
            "description": "Confirm the subscreen opens on Summary tab by default.",
            "priority": "P2",
            "type": "UI/Functional",
            "precondition": "A configuration exists.",
            "steps": [
                {"action": "Open View for a configuration", "expected": "Summary tab is selected by default"},
            ],
            "postcondition": "Summary tab visible.",
        },
        {
            "id": "TC_ClientPortalConfig_View_006",
            "objective": "Display of core summary fields",
            "description": "Validate values shown for Configuration Name, Login Type, Save VOB, Active.",
            "priority": "P1",
            "type": "Functional",
            "precondition": "Known configuration with predetermined values exists.",
            "steps": [
                {"action": "Open View for the known configuration", "expected": "Configuration Name displays the exact saved name"},
                {"action": "Verify Login Type", "expected": "Displays stored login type (e.g., MFA/SSO/Password)"},
                {"action": "Verify Save VOB and Active", "expected": "Checkboxes reflect stored states accurately"},
            ],
            "postcondition": "Correct values displayed.",
        },
        {
            "id": "TC_ClientPortalConfig_View_007",
            "objective": "Display of feature flags on Summary",
            "description": "Validate Patient Search, Current Activity, Client Authorization, Multi Role flags.",
            "priority": "P1",
            "type": "Functional",
            "precondition": "Configuration with a known combination of flags exists.",
            "steps": [
                {"action": "Open View for the target configuration", "expected": "All four flags reflect exact stored states (checked/unchecked)"},
            ],
            "postcondition": "Flags correctly shown.",
        },
        {
            "id": "TC_ClientPortalConfig_View_008",
            "objective": "User Role display",
            "description": "Ensure the User Role field shows the saved role.",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "Configuration with a known role exists.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "User Role displays the stored role label; field is read-only"},
            ],
            "postcondition": "Role correctly displayed.",
        },
        {
            "id": "TC_ClientPortalConfig_View_009",
            "objective": "Provider Search selection display",
            "description": "Verify the Provider Search list shows the saved option(s).",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "Configuration has a known Provider Search selection.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "Provider Search list highlights the saved option(s)"},
            ],
            "postcondition": "Provider Search reflects stored value.",
        },
        {
            "id": "TC_ClientPortalConfig_View_010",
            "objective": "Reports list content accuracy",
            "description": "Validate all assigned reports are displayed with no duplicates.",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "Configuration has a known set of assigned reports.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "All assigned reports are present; order is consistent; no duplicates"},
                {"action": "Scroll through the list (if long)", "expected": "Scrollbar works and all items become visible"},
            ],
            "postcondition": "Reports list correctly rendered.",
        },
        {
            "id": "TC_ClientPortalConfig_View_011",
            "objective": "Account Managers list accuracy",
            "description": "Validate assigned account managers display correctly.",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "Configuration has known assigned account manager(s).",
            "steps": [
                {"action": "Open View for the configuration", "expected": "Assigned account managers are listed without duplication"},
            ],
            "postcondition": "Accurate manager list.",
        },
        {
            "id": "TC_ClientPortalConfig_View_012",
            "objective": "Clients list accuracy and scrolling",
            "description": "Validate all mapped clients are listed and scroll works.",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "Configuration mapped to multiple clients.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "All mapped clients are shown; no duplicates"},
                {"action": "Scroll to bottom", "expected": "All clients can be viewed via scrollbar"},
            ],
            "postcondition": "Clients list fully viewable.",
        },
        {
            "id": "TC_ClientPortalConfig_View_013",
            "objective": "Display for configurations with no assignments",
            "description": "View a configuration with zero reports/clients/account managers.",
            "priority": "P2",
            "type": "Functional/Edge Case",
            "precondition": "Configuration exists with no assignments.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "Empty lists show as blank or 'No records'; no errors"},
            ],
            "postcondition": "Subscreen stable with empty data.",
        },
        {
            "id": "TC_ClientPortalConfig_View_014",
            "objective": "Special characters and long name rendering",
            "description": "Validate UI handles long names and special characters in Configuration Name.",
            "priority": "P3",
            "type": "UI/Functional",
            "precondition": "Configuration named like 'Client_Config_#a-12345678901234567890'.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "Name displays correctly; truncation/tooltip behaves gracefully; no encoding issues"},
            ],
            "postcondition": "UI renders uncommon characters safely.",
        },
        {
            "id": "TC_ClientPortalConfig_View_015",
            "objective": "Inactive configuration visual state",
            "description": "Verify display when configuration is inactive.",
            "priority": "P2",
            "type": "Functional/UI",
            "precondition": "An inactive configuration exists.",
            "steps": [
                {"action": "Open View for inactive configuration", "expected": "Active is unchecked; any inactive indicators/styles appear as per spec"},
            ],
            "postcondition": "Correct inactive state shown.",
        },
        {
            "id": "TC_ClientPortalConfig_View_016",
            "objective": "Multi Role and User Role consistency",
            "description": "Verify role display when Multi Role is enabled.",
            "priority": "P3",
            "type": "Functional",
            "precondition": "Configuration with Multi Role enabled.",
            "steps": [
                {"action": "Open View for the configuration", "expected": "Multi Role shows checked; User Role displays the primary or designated role label"},
            ],
            "postcondition": "Role information coherent.",
        },
        {
            "id": "TC_ClientPortalConfig_View_017",
            "objective": "Close View using Exit",
            "description": "Verify Exit closes the subscreen and returns to the list.",
            "priority": "P1",
            "type": "Functional/UI",
            "precondition": "View subscreen is open.",
            "steps": [
                {"action": "Click Exit", "expected": "Subscreen closes; user returned to the configuration list; no prompts about unsaved changes"},
            ],
            "postcondition": "Back on list; no changes saved.",
        },
        {
            "id": "TC_ClientPortalConfig_View_018",
            "objective": "Access with view-only permission",
            "description": "Ensure a view-only user can open View and cannot see edit actions.",
            "priority": "P1",
            "type": "Security/Functional",
            "precondition": "User with view-only role exists.",
            "steps": [
                {"action": "Login as view-only user and open the list", "expected": "List loads successfully"},
                {"action": "Click View on a configuration", "expected": "View opens read-only; no edit/save/delete controls visible"},
            ],
            "postcondition": "Access limited to viewing.",
        },
        {
            "id": "TC_ClientPortalConfig_View_019",
            "objective": "Block access for unauthorized users",
            "description": "Ensure users without access cannot open View.",
            "priority": "P1",
            "type": "Security/Negative",
            "precondition": "User without module permission exists.",
            "steps": [
                {"action": "Login as unauthorized user, attempt to access configuration list and click View", "expected": "Access denied message or module not visible; View cannot be opened"},
            ],
            "postcondition": "No unauthorized access.",
        },
        {
            "id": "TC_ClientPortalConfig_View_020",
            "objective": "Performance of opening View",
            "description": "Measure load time of the View subscreen.",
            "priority": "P3",
            "type": "Performance",
            "precondition": "Average network/app load.",
            "steps": [
                {"action": "Click View on a configuration and time the load", "expected": "View subscreen renders within <=3 seconds"},
            ],
            "postcondition": "Performance within threshold.",
        },
        {
            "id": "TC_ClientPortalConfig_View_021",
            "objective": "No audit/last-modified impact on view",
            "description": "Ensure viewing does not update audit trails or modified timestamps.",
            "priority": "P3",
            "type": "Functional/Non-Functional",
            "precondition": "Record has known Last Updated timestamp; audit trail accessible.",
            "steps": [
                {"action": "Note current Last Updated/Audit info", "expected": "Timestamp and audit entries recorded"},
                {"action": "Open and close View", "expected": "No new audit entry for 'view'; Last Updated unchanged"},
            ],
            "postcondition": "Metadata intact.",
        },
        {
            "id": "TC_ClientPortalConfig_View_022",
            "objective": "Backend consistency of displayed data",
            "description": "Validate all values match database/service response.",
            "priority": "P2",
            "type": "Functional/Data Integrity",
            "precondition": "Ability to query DB or API.",
            "steps": [
                {"action": "Open View for a known configuration", "expected": "Displayed values captured (fields, flags, lists)"},
                {"action": "Query DB/API for the same configuration", "expected": "All UI values match backend; no missing or extra items"},
            ],
            "postcondition": "UI and backend consistent.",
        },
        {
            "id": "TC_ClientPortalConfig_View_023",
            "objective": "Stability with very long lists",
            "description": "Validate View handles large client/report lists without UI issues.",
            "priority": "P2",
            "type": "Functional/Usability",
            "precondition": "Configuration mapped to 200+ clients and 50+ reports.",
            "steps": [
                {"action": "Open View for the heavy configuration", "expected": "Subscreen loads successfully; scrolling remains smooth; no truncation beyond control capacity"},
            ],
            "postcondition": "Large lists displayed reliably.",
        },
        {
            "id": "TC_ClientPortalConfig_View_024",
            "objective": "Cross-check with Edit view",
            "description": "Ensure values in View match those in Edit.",
            "priority": "P2",
            "type": "Functional",
            "precondition": "User has edit permission.",
            "steps": [
                {"action": "Open View for a configuration and note key values", "expected": "Values recorded"},
                {"action": "Close View, open the same record in Edit", "expected": "All values match between View and Edit"},
            ],
            "postcondition": "Consistency confirmed.",
        },
    ]

    # Merge defaults and add to document
    for tc in tcs:
        tc_with_defaults = {**default_meta, **tc}
        add_test_case(document, tc_with_defaults)

    document.save(output_path)


if __name__ == "__main__":
    OUTPUT = "/workspace/ClientPortalConfig_View_TestCases.docx"
    build_document(OUTPUT)
    print(f"Generated: {OUTPUT}")
