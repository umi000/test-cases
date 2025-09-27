from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_metadata_table(document: Document, meta: dict) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for key in [
        'Test Case ID', 'Test Objective', 'Description', 'Module', 'Priority',
        'Test Type', 'Date', 'Tester', 'JIRA issue #', 'Precondition'
    ]:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = meta.get(key, '')


def add_steps_table(document: Document, steps: list) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step #'
    hdr_cells[1].text = 'User Action'
    hdr_cells[2].text = 'Expected Result'
    hdr_cells[3].text = 'Actual Result'

    for step in steps:
        row_cells = table.add_row().cells
        row_cells[0].text = str(step.get('step', ''))
        row_cells[1].text = step.get('action', '')
        row_cells[2].text = step.get('expected', '')
        row_cells[3].text = step.get('actual', '[Tester to fill after execution]')


def add_post_condition(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(f"Post Condition: {text}")
    run.font.size = Pt(10)


def add_test_case(document: Document, tc: dict) -> None:
    title = document.add_heading(level=2)
    title_run = title.add_run(f"{tc['Test Case ID']} – {tc['Test Objective']}")
    title_run.bold = True

    add_metadata_table(document, tc)
    document.add_paragraph()  # spacer

    # Steps heading
    sh = document.add_paragraph()
    sh_run = sh.add_run('Test Steps:')
    sh_run.bold = True

    add_steps_table(document, tc['steps'])
    document.add_paragraph()  # spacer
    add_post_condition(document, tc.get('Post Condition', ''))
    document.add_page_break()


def build_test_cases() -> list:
    tester = 'Umair Aslam'
    date = '09/25/2025'

    cases = []

    # Helper for creating a case quickly
    def c(
        _id, objective, description, module, priority, test_type,
        precondition, steps, post_condition
    ):
        return {
            'Test Case ID': _id,
            'Test Objective': objective,
            'Description': description,
            'Module': module,
            'Priority': priority,
            'Test Type': test_type,
            'Date': date,
            'Tester': tester,
            'JIRA issue #': '[Insert JIRA ID]',
            'Precondition': precondition,
            'steps': steps,
            'Post Condition': post_condition,
        }

    module = 'Claim Fee Maintenance'

    # 001
    cases.append(c(
        'TC_CFM_001',
        'Verify that the Claim Fee Maintenance screen loads successfully.',
        'Ensure that the screen is displayed when accessed via menu.',
        module, 'P1', 'Functional Test Case',
        'User has access to Claim Fee Maintenance.',
        [
            {'step': 1, 'action': 'Navigate System > Utilities > Claim Fee Type Maintenance',
             'expected': 'Claim Fee Maintenance screen is displayed'},
        ],
        'User is on the Claim Fee Maintenance screen.'
    ))

    # 002
    cases.append(c(
        'TC_CFM_002', 'Verify grid shows exactly five required columns.',
        'Confirm headers Seq, Description, Description 2, Amount, Currency.',
        module, 'P1', 'UI/Functional',
        'Claim Fee Maintenance screen is open.',
        [
            {'step': 1, 'action': 'Select any client and click View',
             'expected': 'Grid renders with exactly 5 columns: Seq, Description, Description 2, Amount, Currency'},
        ],
        'Grid displayed with correct columns.'
    ))

    # 003
    cases.append(c(
        'TC_CFM_003', 'Verify client selection populates fee records.',
        'Selecting a client loads that client’s fees.',
        module, 'P1', 'Functional',
        'At least one client with existing fees.',
        [
            {'step': 1, 'action': 'Choose Client B and click View',
             'expected': 'Grid displays Client B’s fee records in DB sort order'},
        ],
        'Fees for selected client are visible.'
    ))

    # 004
    cases.append(c(
        'TC_CFM_004', 'Verify empty-state rendering for a client without fees.',
        'Ensure proper message and action availability.',
        module, 'P2', 'UI/Functional',
        'Client A exists with no fee records.',
        [
            {'step': 1, 'action': 'Select Client A and click View',
             'expected': 'Grid shows no data message; Add and Copy From are enabled'},
        ],
        'Empty grid with actions available.'
    ))

    # 005
    cases.append(c(
        'TC_CFM_005', 'Verify Amount displays two decimals (XXX.00).',
        'Ensure formatting in grid.',
        module, 'P1', 'UI Validation',
        'Client with fees having various amounts.',
        [
            {'step': 1, 'action': 'View client with fees',
             'expected': 'Amount column values show exactly two decimals (e.g., 15.50)'},
        ],
        'Amount formatting verified.'
    ))

    # 006
    cases.append(c(
        'TC_CFM_006', 'Verify Currency shows ISO code only (e.g., USD).',
        'Ensure display uses currency code, not symbol.',
        module, 'P2', 'UI Validation',
        'Client with mixed currencies.',
        [
            {'step': 1, 'action': 'View client with fees',
             'expected': 'Currency column displays codes (USD, EUR, JPY), no symbols'},
        ],
        'Currency display verified.'
    ))

    # 010
    cases.append(c(
        'TC_CFM_010', 'Verify Add form opens with required fields.',
        'Validate presence of Description, Description 2, Amount, Currency, OK, Cancel.',
        module, 'P1', 'Functional',
        'Grid loaded.',
        [
            {'step': 1, 'action': 'Click Add',
             'expected': 'Add dialog opens with the listed fields and buttons'},
        ],
        'Add dialog visible.'
    ))

    # 011
    cases.append(c(
        'TC_CFM_011', 'Verify default values in Add form.',
        'Amount defaults to 0.00; Currency defaults to USD.',
        module, 'P1', 'Validation',
        'Add form open.',
        [
            {'step': 1, 'action': 'Observe Amount and Currency defaults',
             'expected': 'Amount = 0.00; Currency = USD'},
        ],
        'Defaults confirmed.'
    ))

    # 012
    cases.append(c(
        'TC_CFM_012', 'Save a new fee with minimal required fields.',
        'Descriptions empty; Amount=0.00; Currency=USD.',
        module, 'P1', 'Functional',
        'Add form open.',
        [
            {'step': 1, 'action': 'Leave descriptions blank',
             'expected': 'No client-side error'},
            {'step': 2, 'action': 'Click OK',
             'expected': 'Record saved; appears in grid with auto-generated Seq; Amount 0.00; Currency USD'},
        ],
        'New row saved.'
    ))

    # 013
    cases.append(c(
        'TC_CFM_013', 'Save a new fee with full data including Unicode.',
        'Verify acceptance of all characters in descriptions.',
        module, 'P2', 'Functional/Validation',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Description with Unicode (e.g., “Expedítèd 📦”)',
             'expected': 'Field accepts input'},
            {'step': 2, 'action': 'Set Amount 10.50, Currency EUR; click OK',
             'expected': 'Record saved; grid shows entered values'},
        ],
        'Row saved with Unicode content.'
    ))

    # 014
    cases.append(c(
        'TC_CFM_014', 'Validate Amount is required.',
        'Missing Amount should block save on submit.',
        module, 'P1', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Clear Amount and click OK',
             'expected': 'Error message shown; form not saved'},
        ],
        'No record created.'
    ))

    # 015
    cases.append(c(
        'TC_CFM_015', 'Reject non-numeric Amount.',
        'Letters/specials not allowed.',
        module, 'P1', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Amount “abc” and click OK',
             'expected': 'Validation error; no save'},
        ],
        'No record created.'
    ))

    # 016
    cases.append(c(
        'TC_CFM_016', 'Enforce two-decimal precision.',
        'Amount with >2 decimals blocked.',
        module, 'P1', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Amount 10.123; click OK',
             'expected': 'Error indicating two-decimal format required; no save'},
        ],
        'No record created.'
    ))

    # 017
    cases.append(c(
        'TC_CFM_017', 'Reject negative Amount values.',
        'Amount must be >= 0.00.',
        module, 'P1', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Amount -1.00; click OK',
             'expected': 'Error message; form not saved'},
        ],
        'No record created.'
    ))

    # 018
    cases.append(c(
        'TC_CFM_018', 'Accept zero Amount.',
        '0.00 is valid.',
        module, 'P2', 'Validation Edge',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Amount 0.00; click OK',
             'expected': 'Record saved; grid shows 0.00'},
        ],
        'Row saved.'
    ))

    # 019
    cases.append(c(
        'TC_CFM_019', 'Validate large Amount within limits.',
        'Accept max allowed; reject above.',
        module, 'P2', 'Validation Edge',
        'Add form open; know system max.',
        [
            {'step': 1, 'action': 'Enter max allowed (e.g., 999999999999.99); click OK',
             'expected': 'Save succeeds'},
            {'step': 2, 'action': 'Try max + 0.01',
             'expected': 'Validation error; no save'},
        ],
        'Only valid max saved.'
    ))

    # 020
    cases.append(c(
        'TC_CFM_020', 'Reject Amount with thousand separators if not supported.',
        '“1,234.00” should be invalid.',
        module, 'P2', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Enter Amount 1,234.00; click OK',
             'expected': 'Error shown; no save'},
        ],
        'No record created.'
    ))

    # 021
    cases.append(c(
        'TC_CFM_021', 'Validate Currency selection required.',
        'Missing/cleared currency blocked.',
        module, 'P1', 'Validation Negative',
        'Add form open.',
        [
            {'step': 1, 'action': 'Clear Currency (if possible) and click OK',
             'expected': 'Error near Currency; no save'},
        ],
        'No record created.'
    ))

    # 022
    cases.append(c(
        'TC_CFM_022', 'Prevent multiple creations on rapid double-submit.',
        'Debounce duplicate save clicks/Enter.',
        module, 'P1', 'Functional/Robustness',
        'Add form with valid data.',
        [
            {'step': 1, 'action': 'Double-click OK rapidly (or press Enter repeatedly)',
             'expected': 'Exactly one record is created'},
        ],
        'Single row created.'
    ))

    # 023
    cases.append(c(
        'TC_CFM_023', 'Cancel in Add discards changes.',
        'No record created on cancel.',
        module, 'P2', 'UI/Functional',
        'Add form with unsaved data.',
        [
            {'step': 1, 'action': 'Click Cancel',
             'expected': 'Dialog closes; grid unchanged'},
        ],
        'No changes.'
    ))

    # 030
    cases.append(c(
        'TC_CFM_030', 'Edit form opens pre-filled.',
        'Ensure values load correctly.',
        module, 'P1', 'Functional',
        'At least one existing row.',
        [
            {'step': 1, 'action': 'Click Edit on a row',
             'expected': 'Edit dialog opens with existing values populated'},
        ],
        'Edit dialog visible.'
    ))

    # 031
    cases.append(c(
        'TC_CFM_031', 'Edit descriptions only.',
        'Save updates without changing amount/currency.',
        module, 'P2', 'Functional',
        'Edit dialog open.',
        [
            {'step': 1, 'action': 'Change Description fields; click OK',
             'expected': 'Row updated; amount/currency unchanged'},
        ],
        'Row updated.'
    ))

    # 032
    cases.append(c(
        'TC_CFM_032', 'Edit Amount to valid two-decimal value.',
        'Save success.',
        module, 'P1', 'Functional/Validation',
        'Edit dialog open.',
        [
            {'step': 1, 'action': 'Set Amount to 25.00; click OK',
             'expected': 'Row saved; grid shows 25.00'},
        ],
        'Row updated.'
    ))

    # 033
    cases.append(c(
        'TC_CFM_033', 'Edit validation – clear Amount rejected.',
        'Amount required on submit.',
        module, 'P1', 'Validation Negative',
        'Edit dialog open.',
        [
            {'step': 1, 'action': 'Clear Amount and click OK',
             'expected': 'Error shown; no save'},
        ],
        'No change.'
    ))

    # 034
    cases.append(c(
        'TC_CFM_034', 'Edit validation – negative Amount rejected.',
        'Must be >= 0.00.',
        module, 'P1', 'Validation Negative',
        'Edit dialog open.',
        [
            {'step': 1, 'action': 'Enter -5.00; click OK',
             'expected': 'Error shown; no save'},
        ],
        'No change.'
    ))

    # 035
    cases.append(c(
        'TC_CFM_035', 'Edit validation – more than two decimals rejected.',
        'Precision enforcement.',
        module, 'P1', 'Validation Negative',
        'Edit dialog open.',
        [
            {'step': 1, 'action': 'Enter 10.123; click OK',
             'expected': 'Error shown; no save'},
        ],
        'No change.'
    ))

    # 036
    cases.append(c(
        'TC_CFM_036', 'Debounce duplicate edit saves.',
        'Prevent multiple updates on rapid clicks.',
        module, 'P2', 'Functional/Robustness',
        'Edit dialog with valid changes.',
        [
            {'step': 1, 'action': 'Double-click OK rapidly',
             'expected': 'Only one update recorded'},
        ],
        'Single update.'
    ))

    # 037
    cases.append(c(
        'TC_CFM_037', 'Cancel in Edit discards changes.',
        'Ensure no persistence on cancel.',
        module, 'P2', 'UI/Functional',
        'Edit dialog with unsaved changes.',
        [
            {'step': 1, 'action': 'Click Cancel',
             'expected': 'Dialog closes; row remains unchanged'},
        ],
        'No changes.'
    ))

    # 040
    cases.append(c(
        'TC_CFM_040', 'Verify Delete confirmation dialog appears.',
        'Confirm destructive action prompt text.',
        module, 'P1', 'Functional/UI',
        'At least one row exists.',
        [
            {'step': 1, 'action': 'Click Delete on a row',
             'expected': 'Dialog shows: “Are you sure you want to delete this record?” with Yes/No (or OK/Cancel)'},
        ],
        'Confirmation shown.'
    ))

    # 041
    cases.append(c(
        'TC_CFM_041', 'Confirming delete removes record.',
        'Verify persistence after refresh.',
        module, 'P1', 'Functional',
        'Delete dialog open.',
        [
            {'step': 1, 'action': 'Click Yes/OK',
             'expected': 'Row removed from grid'},
            {'step': 2, 'action': 'Refresh or re-open client',
             'expected': 'Row remains deleted'},
        ],
        'Row deleted.'
    ))

    # 042
    cases.append(c(
        'TC_CFM_042', 'Cancel delete leaves record intact.',
        'Validate no change on cancel.',
        module, 'P2', 'Functional/UI',
        'Delete dialog open.',
        [
            {'step': 1, 'action': 'Click No/Cancel',
             'expected': 'Dialog closes; row remains'},
        ],
        'No changes.'
    ))

    # 050
    cases.append(c(
        'TC_CFM_050', 'Reorder – move item up and save.',
        'Verify Move Up updates order on save.',
        module, 'P1', 'Functional',
        'Client with at least 3 rows; “No Fee Selected” row (if present) is immovable.',
        [
            {'step': 1, 'action': 'Click Reorder',
             'expected': 'Reorder mode opens with Move Up/Down controls'},
            {'step': 2, 'action': 'Select a middle row; click Move Up once',
             'expected': 'Row visually moves up by one'},
            {'step': 3, 'action': 'Click Save',
             'expected': 'Grid reflects new order'},
        ],
        'Order persisted.'
    ))

    # 051
    cases.append(c(
        'TC_CFM_051', 'Reorder – move item down and save.',
        'Verify Move Down.',
        module, 'P1', 'Functional',
        'Client with at least 3 rows.',
        [
            {'step': 1, 'action': 'Enter Reorder; select a middle row; click Move Down',
             'expected': 'Row moves down by one'},
            {'step': 2, 'action': 'Click Save',
             'expected': 'New order saved'},
        ],
        'Order persisted.'
    ))

    # 052
    cases.append(c(
        'TC_CFM_052', 'Reorder – Cancel discards changes.',
        'Ensure no reordering persisted on cancel.',
        module, 'P2', 'UI/Functional',
        'Reorder mode with unsaved moves.',
        [
            {'step': 1, 'action': 'Click Cancel',
             'expected': 'Reorder closes; original order restored'},
        ],
        'Original order remains.'
    ))

    # 053
    cases.append(c(
        'TC_CFM_053', 'Reorder – boundary controls disabled.',
        'First movable item cannot move up; last cannot move down.',
        module, 'P2', 'UI Validation',
        'Reorder mode open.',
        [
            {'step': 1, 'action': 'Select first movable row',
             'expected': 'Move Up disabled/inactive'},
            {'step': 2, 'action': 'Select last row',
             'expected': 'Move Down disabled/inactive'},
        ],
        'Controls validated.'
    ))

    # 054
    cases.append(c(
        'TC_CFM_054', 'Reorder – “No Fee Selected” row cannot move.',
        'Protected row immovable.',
        module, 'P2', 'Functional/UI',
        'Protected row exists.',
        [
            {'step': 1, 'action': 'Try Move Up/Down on “No Fee Selected”',
             'expected': 'Action blocked; controls disabled'},
        ],
        'Protected row position unchanged.'
    ))

    # 060
    cases.append(c(
        'TC_CFM_060', 'Copy From – open selection UI.',
        'Verify Copy From dialog shows source selector.',
        module, 'P2', 'Functional/UI',
        'Screen open.',
        [
            {'step': 1, 'action': 'Click Copy From',
             'expected': 'Dialog opens to choose source client/template'},
        ],
        'Copy From dialog visible.'
    ))

    # 061
    cases.append(c(
        'TC_CFM_061', 'Copy From into empty client.',
        'Ensure records are appended to empty list.',
        module, 'P2', 'Functional',
        'Current client has zero fees; at least one source exists.',
        [
            {'step': 1, 'action': 'Click Copy From; select a source; confirm',
             'expected': 'Fees copied; amounts/currencies retained; new Seq assigned'},
        ],
        'Destination populated.'
    ))

    # 062
    cases.append(c(
        'TC_CFM_062', 'Copy From into non-empty client (append).',
        'Verify existing records remain; source records added.',
        module, 'P2', 'Functional',
        'Destination has fees.',
        [
            {'step': 1, 'action': 'Perform Copy From as above',
             'expected': 'Existing rows retained; source rows appended as per design'},
        ],
        'List contains union of rows.'
    ))

    # 063
    cases.append(c(
        'TC_CFM_063', 'Cancel Copy From makes no changes.',
        'Ensure cancel behavior.',
        module, 'P3', 'UI/Functional',
        'Copy From dialog open.',
        [
            {'step': 1, 'action': 'Click Cancel',
             'expected': 'Dialog closes; grid unchanged'},
        ],
        'No changes.'
    ))

    # 070
    cases.append(c(
        'TC_CFM_070', 'Amount whitespace trimmed before validation.',
        '“ 10.00 ” treated as 10.00 and valid.',
        module, 'P3', 'Validation Edge',
        'Add or Edit form.',
        [
            {'step': 1, 'action': 'Enter Amount with spaces; submit',
             'expected': 'Save succeeds; stored/displayed as 10.00'},
        ],
        'Value normalized.'
    ))

    # 071
    cases.append(c(
        'TC_CFM_071', 'Reject comma decimal formats.',
        '“10,50” invalid; suggest dot separator.',
        module, 'P2', 'Validation Negative',
        'Add/Edit form.',
        [
            {'step': 1, 'action': 'Enter 10,50; submit',
             'expected': 'Error indicates dot-decimal required; no save'},
        ],
        'No save.'
    ))

    # 072
    cases.append(c(
        'TC_CFM_072', 'Reject scientific notation for Amount.',
        '“1e3” invalid.',
        module, 'P3', 'Validation Negative',
        'Add/Edit form.',
        [
            {'step': 1, 'action': 'Enter 1e3; submit',
             'expected': 'Validation error; no save'},
        ],
        'No save.'
    ))

    # 073
    cases.append(c(
        'TC_CFM_073', 'Accept extremely long descriptions without UI break.',
        'Save and verify grid truncation/ellipsis.',
        module, 'P3', 'UI/Functional Edge',
        'Add/Edit form.',
        [
            {'step': 1, 'action': 'Paste 10,000 chars in Description; valid Amount/Currency; save',
             'expected': 'Save succeeds; grid displays without layout break (may truncate visually)'},
        ],
        'Row saved; UI stable.'
    ))

    # 080
    cases.append(c(
        'TC_CFM_080', 'Seq assigned on create.',
        'New record gets next Seq value.',
        module, 'P2', 'Functional',
        'Note highest Seq; create new row.',
        [
            {'step': 1, 'action': 'Add valid record',
             'expected': 'New row displays with appropriate next Seq'},
        ],
        'Seq validated.'
    ))

    # 081
    cases.append(c(
        'TC_CFM_081', 'Seq reflects new order after reorder.',
        'Ensure uniqueness and order post-save.',
        module, 'P2', 'Functional',
        'Perform reorder and save.',
        [
            {'step': 1, 'action': 'Inspect Seq values',
             'expected': 'Seqs reflect reordered list and remain unique'},
        ],
        'Seq consistent.'
    ))

    # 082
    cases.append(c(
        'TC_CFM_082', 'Delete does not corrupt Seq/order upon reload.',
        'Verify stability after deletion.',
        module, 'P3', 'Functional',
        'Existing multiple rows.',
        [
            {'step': 1, 'action': 'Delete a middle row; reload client',
             'expected': 'Remaining rows display; order consistent with DB'},
        ],
        'Grid stable.'
    ))

    # 090
    cases.append(c(
        'TC_CFM_090', 'Page load performance with 0 records.',
        'Validate SLA (e.g., <1s local).',
        module, 'P2', 'Non-Functional Performance',
        'Empty client.',
        [
            {'step': 1, 'action': 'Load screen; select empty client; View',
             'expected': 'Page and grid load within SLA; no console errors'},
        ],
        'Performance recorded.'
    ))

    # 091
    cases.append(c(
        'TC_CFM_091', 'Page load performance with large dataset.',
        'Validate SLA with 500+ rows.',
        module, 'P2', 'Non-Functional Performance',
        'Client C with ≥500 fees.',
        [
            {'step': 1, 'action': 'Select Client C; click View',
             'expected': 'Loads within SLA; UI responsive; no memory spikes'},
        ],
        'Performance acceptable.'
    ))

    # 092
    cases.append(c(
        'TC_CFM_092', 'Operation latency for Add/Edit/Delete/Reorder.',
        'Each server response under agreed SLA (e.g., <1s).',
        module, 'P3', 'Non-Functional Performance',
        'Network stable.',
        [
            {'step': 1, 'action': 'Perform Add/Edit/Delete/Reorder',
             'expected': 'Each completes within SLA; success toast shown'},
        ],
        'Latency recorded.'
    ))

    # 093
    cases.append(c(
        'TC_CFM_093', 'Keyboard accessibility and focus management.',
        'Ensure tab order, visible focus, keyboard reorder.',
        module, 'P2', 'Accessibility',
        'Screen open.',
        [
            {'step': 1, 'action': 'Navigate with Tab/Shift+Tab across controls',
             'expected': 'All focusable elements reachable with visible focus'},
            {'step': 2, 'action': 'Open Reorder; use keyboard to move items',
             'expected': 'Items move via buttons/shortcuts; screen reader announces changes'},
        ],
        'A11Y validated.'
    ))

    # 094
    cases.append(c(
        'TC_CFM_094', 'Dialog ARIA and announcements.',
        'Add/Edit/Delete dialogs announce open/close and errors.',
        module, 'P3', 'Accessibility',
        'Screen reader active.',
        [
            {'step': 1, 'action': 'Open Add/Edit/Delete dialog',
             'expected': 'Dialog has role="dialog"; focus trapped; announced'},
            {'step': 2, 'action': 'Trigger validation error',
             'expected': 'Error announced and associated to field'},
        ],
        'A11Y announcements verified.'
    ))

    # 095
    cases.append(c(
        'TC_CFM_095', 'Cross-browser compatibility.',
        'Validate core flows on Chrome, Firefox, Edge, Safari.',
        module, 'P2', 'Cross-Browser',
        'Browsers installed.',
        [
            {'step': 1, 'action': 'Repeat core flows (View/Add/Edit/Delete/Reorder/Copy) in each browser',
             'expected': 'Identical behavior; consistent formatting; no browser-specific errors'},
        ],
        'Cross-browser status captured.'
    ))

    # 096
    cases.append(c(
        'TC_CFM_096', 'Security – escape HTML in descriptions.',
        'Ensure stored output is not executed.',
        module, 'P3', 'Security Negative',
        'Add/Edit form.',
        [
            {'step': 1, 'action': 'Enter Description as <script>alert(1)</script>; valid Amount/Currency; save',
             'expected': 'Value saved and displayed as text; no script execution'},
        ],
        'Output safely escaped.'
    ))

    return cases


def main(output_path: str) -> None:
    document = Document()

    # Title
    title = document.add_heading('Claim Fee Maintenance – Test Cases', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph('All test cases exported with steps in tabular format.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tc in build_test_cases():
        add_test_case(document, tc)

    document.save(output_path)


if __name__ == '__main__':
    # Default output location
    output = '/workspace/Claim_Fee_Maintenance_Test_Cases.docx'
    main(output)

